#!/usr/bin/env python3
"""
3GPP Word document structure extractor.

Downloads a ZIP from 3GPP FTP, extracts the .docx/.doc files, and converts
them to the same _structure.jsonl format that extract_pdf_structure.py produces.

Supports:
  - .docx (modern OOXML, via python-docx)
  - .doc  (old binary format, via antiword command-line tool)

Usage:
    python scripts/extract_docx_structure.py <spec_id> <3gpp_ftp_url>
    python scripts/extract_docx_structure.py ts_34_229_1 \
        https://www.3gpp.org/ftp/Specs/archive/34_series/34.229-1/34229-1-j50.zip

Output:
    data/intermediate/{spec_id}_structure.jsonl
"""
from __future__ import annotations

import argparse
import io
import json
import os
import re
import subprocess
import sys
import tempfile
import urllib.request
import zipfile
from pathlib import Path

try:
    import docx
except ImportError:
    print("ERROR: python-docx required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)

_scripts_dir = Path(__file__).resolve().parent
if str(_scripts_dir) not in sys.path:
    sys.path.insert(0, str(_scripts_dir))

from common_paths import INTERMEDIATE_DIR

USER_AGENT = "Mozilla/5.0 (compatible; 3GPP-Downloader/1.0)"

# 3GPP heading style names
HEADING_STYLES = {
    "heading 1", "heading 2", "heading 3", "heading 4",
    "heading 5", "heading 6", "heading 7", "heading 8",
    "heading1", "heading2", "heading3", "heading4",
    # 3GPP-specific style names sometimes seen
    "3gpp heading 1", "3gpp heading 2",
}

SECTION_PATTERN = re.compile(r'^(\d+(?:\.\d+)*)\s+(.+)', re.DOTALL)
ANNEX_PATTERN = re.compile(
    r'^(Annex\s+[A-Z](?:\.\d+)*)\s*(?:\((?:normative|informative)\))?\s*:?\s*(.+)',
    re.DOTALL,
)

TOP_LEVEL_HEADINGS = {
    'foreword', 'introduction', 'scope', 'references',
    'definitions', 'abbreviations',
}


def _is_heading_style(style_name: str) -> bool:
    return style_name.lower() in HEADING_STYLES


def _heading_level(style_name: str) -> int:
    m = re.search(r'(\d+)', style_name)
    return int(m.group(1)) if m else 1


def _extract_section_number(text: str) -> str:
    m = SECTION_PATTERN.match(text.strip())
    if m:
        return m.group(1)
    m = ANNEX_PATTERN.match(text.strip())
    if m:
        return m.group(1)
    return ""


def download_zip(url: str) -> bytes:
    """Download a ZIP file and return raw bytes."""
    print(f"  Downloading {url}...")
    req = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
    with urllib.request.urlopen(req, timeout=120) as resp:
        data = resp.read()
    print(f"  Downloaded {len(data) / 1024:.0f} KB")
    return data


def extract_docx_from_zip(zip_bytes: bytes) -> tuple[bytes, str]:
    """Extract the first .docx file from a ZIP archive (single-file fallback)."""
    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
        docx_files = [n for n in zf.namelist() if n.lower().endswith(".docx")]
        if not docx_files:
            raise ValueError("No .docx file found in ZIP archive")
        name = docx_files[0]
        return zf.read(name), name


def extract_all_docx_from_zip(zip_bytes: bytes) -> list[tuple[bytes, str]]:
    """Extract ALL .docx/.doc files from a ZIP archive, sorted by name, excluding cover."""
    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
        all_files = zf.namelist()
        # Prefer .docx; fall back to .doc
        docx_files = sorted(
            n for n in all_files
            if n.lower().endswith(".docx") and "cover" not in n.lower()
        )
        if not docx_files:
            docx_files = sorted(
                n for n in all_files
                if n.lower().endswith(".doc") and "cover" not in n.lower()
            )
        if not docx_files:
            # Include cover as last resort
            docx_files = sorted(
                n for n in all_files
                if n.lower().endswith((".docx", ".doc"))
            )
        if not docx_files:
            raise ValueError("No .docx or .doc file found in ZIP archive")
        return [(zf.read(name), name) for name in docx_files]


def extract_text_from_doc(doc_bytes: bytes, filename: str) -> str:
    """Extract text from old binary .doc format using antiword."""
    # antiword needs a real file on disk
    suffix = Path(filename).suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tf:
        tf.write(doc_bytes)
        tf_name = tf.name
    try:
        result = subprocess.run(
            ["antiword", tf_name],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode == 0:
            return result.stdout
        # antiword may fail for encrypted or corrupted files
        return ""
    except (subprocess.TimeoutExpired, FileNotFoundError):
        return ""
    finally:
        try:
            os.unlink(tf_name)
        except OSError:
            pass


def parse_doc_text_to_sections(text: str) -> list[tuple[str, int, str]]:
    """
    Parse antiword text output to find section headings.

    Returns list of (heading_text, level, section_number).
    """
    sections = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        m = SECTION_PATTERN.match(line)
        if m:
            sec_num = m.group(1)
            depth = sec_num.count(".") + 1
            if depth <= 6:  # skip deep subsections as headings
                sections.append((line, depth, sec_num))
            continue
        m2 = ANNEX_PATTERN.match(line)
        if m2:
            sections.append((line, 1, m2.group(1)))
    return sections


def parse_version_from_filename(filename: str) -> str:
    """Try to infer version from 3GPP filename like '34229-1-j50.docx'."""
    # 3GPP naming: NNNNN-P-xYZ.docx or NNNNN-P-xYZ.doc
    # x = letter indicating release (f=15, g=16, h=17, i=18, j=19, k=20)
    # YZ = version digits
    m = re.search(r'-([a-z])(\d{2})\.(docx?|zip)$', filename.lower())
    if m:
        rel_letter = m.group(1)
        ver_digits = m.group(2)
        # Release letter mapping
        letter_to_release = {
            'a': 10, 'b': 11, 'c': 12, 'd': 13, 'e': 14,
            'f': 15, 'g': 16, 'h': 17, 'i': 18, 'j': 19, 'k': 20,
        }
        rel = letter_to_release.get(rel_letter, 0)
        minor = int(ver_digits[0])
        patch = int(ver_digits[1])
        return f"v{rel}.{minor}.{patch}"
    return ""


def extract_docx(
    spec_id: str,
    docx_bytes: bytes,
    source_filename: str,
) -> Path:
    """
    Parse a .docx file and produce a _structure.jsonl in data/intermediate/.

    Returns the output path.
    """
    doc = docx.Document(io.BytesIO(docx_bytes))
    version = parse_version_from_filename(source_filename)

    # Collect paragraphs with style info
    paragraphs = doc.paragraphs

    # Build TOC entries and virtual "pages" from the document body
    # Strategy: each heading starts a new virtual page.
    # Virtual page number = heading index + 1.
    toc_entries: list[dict] = []
    page_texts: list[dict] = []

    current_page = 1
    current_page_lines: list[str] = []
    current_page_headings: list[str] = []

    def flush_page(page_num: int, lines: list[str], headings: list[str]) -> None:
        if lines or headings:
            page_texts.append({
                "type": "page_text",
                "page_number": page_num,
                "text": "\n".join(lines),
                "headings_found": headings[:],
            })

    for para in paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if _is_heading_style(style_name):
            # Flush accumulated content as a page before new heading
            flush_page(current_page, current_page_lines, current_page_headings)
            current_page += 1
            current_page_lines = []
            current_page_headings = []

            level = _heading_level(style_name)
            sec_num = _extract_section_number(text)
            sort_order = len(toc_entries)

            toc_entries.append({
                "type": "toc_entry",
                "source": "docx_style",
                "sort_order": sort_order,
                "level": level,
                "title": text,
                "page": current_page,
                "section_number": sec_num,
            })
            current_page_lines.append(text)
            current_page_headings.append(text)
        else:
            current_page_lines.append(text)

    # Flush remaining content
    flush_page(current_page, current_page_lines, current_page_headings)

    # Fall back: if no heading styles found, scan for 3GPP heading patterns
    if not toc_entries:
        print(f"  [WARN] {spec_id}: No heading styles found, using pattern scan", file=sys.stderr)
        page_texts = []  # reset
        current_page = 1
        current_page_lines = []
        current_page_headings = []

        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue

            m = SECTION_PATTERN.match(text)
            if m:
                flush_page(current_page, current_page_lines, current_page_headings)
                current_page += 1
                current_page_lines = []
                current_page_headings = []

                sec_num = m.group(1)
                depth = sec_num.count(".") + 1
                sort_order = len(toc_entries)
                toc_entries.append({
                    "type": "toc_entry",
                    "source": "pattern_scan",
                    "sort_order": sort_order,
                    "level": depth,
                    "title": text,
                    "page": current_page,
                    "section_number": sec_num,
                })
                current_page_lines.append(text)
                current_page_headings.append(text)
            else:
                m2 = ANNEX_PATTERN.match(text)
                if m2:
                    flush_page(current_page, current_page_lines, current_page_headings)
                    current_page += 1
                    current_page_lines = []
                    current_page_headings = []

                    sort_order = len(toc_entries)
                    toc_entries.append({
                        "type": "toc_entry",
                        "source": "pattern_scan",
                        "sort_order": sort_order,
                        "level": 1,
                        "title": text,
                        "page": current_page,
                        "section_number": m2.group(1),
                    })
                    current_page_lines.append(text)
                    current_page_headings.append(text)
                else:
                    current_page_lines.append(text)

        flush_page(current_page, current_page_lines, current_page_headings)

    total_pages = current_page

    # Build JSONL records
    records: list[dict] = []
    records.append({
        "type": "spec_meta",
        "spec_id": spec_id,
        "title": spec_id,
        "version": version,
        "total_pages": total_pages,
        "source_docx": source_filename,
    })
    records.extend(toc_entries)
    records.extend(page_texts)

    output_path = INTERMEDIATE_DIR / f"{spec_id}_structure.jsonl"
    with open(output_path, "w", encoding="utf-8") as f:
        for record in records:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")

    print(f"  ✓ {spec_id}: {len(toc_entries)} TOC entries, {total_pages} virtual pages → {output_path.name}")
    return output_path


def extract_merged_docx(spec_id: str, docs: list[tuple[bytes, str]]) -> Path:
    """Merge multiple .docx/.doc files and produce a single structure JSONL."""
    all_toc_entries: list[dict] = []
    all_page_texts: list[dict] = []
    page_offset = 0
    version = ""

    for doc_bytes, doc_name in docs:
        print(f"  Processing: {doc_name} ({len(doc_bytes) / 1024:.0f} KB)")
        if not version:
            v = parse_version_from_filename(doc_name)
            if v:
                version = v

        if doc_name.lower().endswith(".doc"):
            # Old binary format: use antiword to extract text
            text = extract_text_from_doc(doc_bytes, doc_name)
            if not text.strip():
                print(f"  [WARN] antiword returned no text for {doc_name}", file=sys.stderr)
                continue
            sections = parse_doc_text_to_sections(text)
            lines = text.split("\n")

            current_page = page_offset + 1
            current_block: list[str] = []

            def flush_block(page_num: int, block: list[str]) -> None:
                if block:
                    all_page_texts.append({
                        "type": "page_text",
                        "page_number": page_num,
                        "text": "\n".join(block),
                        "headings_found": [],
                    })

            for line in lines:
                stripped = line.strip()
                m = SECTION_PATTERN.match(stripped)
                is_heading = bool(m) or bool(ANNEX_PATTERN.match(stripped))
                if is_heading and m:
                    depth = stripped.count(".") + 1 - stripped.count(".") + m.group(1).count(".") + 1
                    depth = m.group(1).count(".") + 1
                    if depth <= 6:
                        flush_block(current_page, current_block)
                        current_page += 1
                        current_block = [stripped]
                        sec_num = m.group(1)
                        all_toc_entries.append({
                            "type": "toc_entry",
                            "source": "antiword_scan",
                            "sort_order": len(all_toc_entries),
                            "level": depth,
                            "title": stripped,
                            "page": current_page,
                            "section_number": sec_num,
                        })
                    else:
                        current_block.append(stripped)
                elif stripped:
                    current_block.append(stripped)

            flush_block(current_page, current_block)
            page_offset = current_page
            continue

        # .docx format: use python-docx
        doc = docx.Document(io.BytesIO(doc_bytes))
        paragraphs = doc.paragraphs

        current_page = page_offset + 1
        current_page_lines: list[str] = []
        current_page_headings: list[str] = []
        had_content = False

        def flush_page_m(page_num: int, lines: list[str], headings: list[str]) -> None:
            if lines or headings:
                all_page_texts.append({
                    "type": "page_text",
                    "page_number": page_num,
                    "text": "\n".join(lines),
                    "headings_found": headings[:],
                })

        for para in paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            had_content = True
            if _is_heading_style(style_name):
                flush_page_m(current_page, current_page_lines, current_page_headings)
                current_page += 1
                current_page_lines = []
                current_page_headings = []

                level = _heading_level(style_name)
                sec_num = _extract_section_number(text)
                sort_order = len(all_toc_entries)

                all_toc_entries.append({
                    "type": "toc_entry",
                    "source": "docx_style",
                    "sort_order": sort_order,
                    "level": level,
                    "title": text,
                    "page": current_page,
                    "section_number": sec_num,
                })
                current_page_lines.append(text)
                current_page_headings.append(text)
            else:
                current_page_lines.append(text)

        flush_page_m(current_page, current_page_lines, current_page_headings)
        if had_content:
            page_offset = current_page

    total_pages = page_offset

    records: list[dict] = []
    records.append({
        "type": "spec_meta",
        "spec_id": spec_id,
        "title": spec_id,
        "version": version,
        "total_pages": total_pages,
        "source_docx": f"{len(docs)}_merged_files",
    })
    records.extend(all_toc_entries)
    records.extend(all_page_texts)

    output_path = INTERMEDIATE_DIR / f"{spec_id}_structure.jsonl"
    with open(output_path, "w", encoding="utf-8") as f:
        for record in records:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")

    print(f"  ✓ {spec_id}: {len(all_toc_entries)} TOC entries, {total_pages} virtual pages → {output_path.name}")
    return output_path


def process(spec_id: str, zip_url: str) -> Path:
    """Download ZIP, extract all .docx files, merge and produce structure JSONL."""
    zip_bytes = download_zip(zip_url)

    all_docs = extract_all_docx_from_zip(zip_bytes)
    print(f"  Found {len(all_docs)} docx file(s) in ZIP")

    if len(all_docs) == 1:
        docx_bytes, docx_name = all_docs[0]
        print(f"  Extracted: {docx_name} ({len(docx_bytes) / 1024:.0f} KB)")
        return extract_docx(spec_id, docx_bytes, docx_name)

    # Multiple docx files: merge them into one structure.jsonl
    return extract_merged_docx(spec_id, all_docs)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Extract 3GPP Word doc structure from 3GPP FTP ZIP.",
        epilog="Output: data/intermediate/{spec_id}_structure.jsonl",
    )
    parser.add_argument("spec_id", help="Spec ID, e.g. ts_34_229_1")
    parser.add_argument("zip_url", help="3GPP FTP ZIP URL")
    args = parser.parse_args()

    process(args.spec_id, args.zip_url)


if __name__ == "__main__":
    main()
